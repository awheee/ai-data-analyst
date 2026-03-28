from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Literal, Tuple

import pandas as pd


@dataclass
class LoadedDataset:
    mode: Literal["dataframe", "text"]
    # When mode == "dataframe"
    df: Optional[pd.DataFrame]
    schema: List[Dict[str, str]]
    # When mode == "text"
    extracted_text: Optional[str]
    # Human-friendly summary shown in UI
    summary: Dict[str, Any]
    # For document modes, keep multiple tables if we extract them
    tables: Optional[List[pd.DataFrame]] = None
    table_index: int = 0


def _normalize_col_name(col: Any) -> str:
    """
    Normalize column names so the LLM can refer to them reliably.
    - lowercase
    - replace non-alnum with underscores
    - collapse repeats
    - trim leading/trailing underscores
    """
    import re

    s = str(col).strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_{2,}", "_", s)
    s = s.strip("_")
    return s or "col"


def _dedupe_columns(cols: List[str]) -> List[str]:
    seen: Dict[str, int] = {}
    out: List[str] = []
    for c in cols:
        if c not in seen:
            seen[c] = 1
            out.append(c)
        else:
            seen[c] += 1
            out.append(f"{c}_{seen[c]}")
    return out


def _dtype_bucket(series: pd.Series) -> str:
    if pd.api.types.is_numeric_dtype(series):
        return "number"
    if pd.api.types.is_datetime64_any_dtype(series):
        return "datetime"
    return "text"


def _build_schema(df: pd.DataFrame) -> List[Dict[str, str]]:
    schema: List[Dict[str, str]] = []
    for col in df.columns:
        schema.append({"name": str(col), "type": _dtype_bucket(df[col])})
    return schema


def _summarize_dataframe(df: pd.DataFrame) -> Dict[str, Any]:
    numeric_cols = [c for c in df.columns if _dtype_bucket(df[c]) == "number"]
    out: Dict[str, Any] = {
        "rows": int(df.shape[0]),
        "columns": int(df.shape[1]),
        "column_names": [str(c) for c in df.columns],
        "numeric_columns": numeric_cols,
    }
    if numeric_cols:
        # Keep output small: describe numeric columns only.
        desc = df[numeric_cols].describe(include="all").transpose()
        out["numeric_describe"] = desc.reset_index(names="column").to_dict(orient="records")
    return out


def build_accurate_column_profile(df: pd.DataFrame, max_top_categorical: int = 12) -> pd.DataFrame:
    """
    Deterministic per-column facts for an accurate dataset summary (no LLM).
    """
    n = len(df)
    rows: List[Dict[str, Any]] = []
    for col in df.columns:
        s = df[col]
        non_null = int(s.notna().sum())
        null_pct = round(100.0 * float(n - non_null) / max(n, 1), 2)
        n_unique = int(s.nunique(dropna=True))
        dtype = str(s.dtype)
        rec: Dict[str, Any] = {
            "column": str(col),
            "dtype": dtype,
            "non_null": non_null,
            "null_pct": null_pct,
            "n_unique": n_unique,
        }
        if pd.api.types.is_numeric_dtype(s):
            if non_null:
                rec["min"] = s.min()
                rec["max"] = s.max()
                rec["mean"] = round(float(s.mean()), 8)
            else:
                rec["min"] = rec["max"] = rec["mean"] = None
        elif pd.api.types.is_datetime64_any_dtype(s):
            rec["min"] = str(s.min()) if non_null else None
            rec["max"] = str(s.max()) if non_null else None
        else:
            if n_unique > 0 and n_unique <= max_top_categorical:
                vc = s.astype(str).value_counts().head(6)
                rec["top_values"] = "; ".join(f"{k} ({v})" for k, v in vc.items())
            else:
                samples = s.dropna().head(3).tolist()
                rec["sample_values"] = "; ".join(str(x)[:100] for x in samples)[:350]
        rows.append(rec)
    return pd.DataFrame(rows)


def build_dataset_narrative(profile_df: pd.DataFrame, n_rows: int, n_cols: int) -> str:
    """
    Short factual readout from the column profile (no LLM). Complements the profile table.
    """
    lines: List[str] = []
    lines.append(f"- **Shape:** {n_rows:,} rows × {n_cols} columns")

    if profile_df.empty:
        return "\n".join(lines)

    dt = profile_df["dtype"].astype(str).str.lower()
    is_dt = dt.str.contains("datetime", na=False)
    is_num = dt.str.contains(r"int|float", regex=True, na=False) & ~is_dt
    n_numeric = int(is_num.sum())
    n_datetime = int(is_dt.sum())
    n_text = max(0, n_cols - n_numeric - n_datetime)
    lines.append(
        f"- **Column types (approx.):** {n_numeric} numeric, {n_datetime} datetime-like, {n_text} other/text"
    )

    high_null = profile_df[profile_df["null_pct"] > 50]
    if not high_null.empty:
        names = ", ".join(f"`{r['column']}` ({r['null_pct']}% null)" for _, r in high_null.head(8).iterrows())
        extra = " …" if len(high_null) > 8 else ""
        lines.append(f"- **High missingness:** {names}{extra}")

    num_rows = profile_df[is_num]
    if not num_rows.empty:
        lines.append(f"- **Numeric columns:** {', '.join('`'+str(c)+'`' for c in num_rows['column'].tolist()[:15])}")
        if len(num_rows) > 15:
            lines.append(f"  - …and {len(num_rows) - 15} more")

    date_rows = profile_df[is_dt]
    if not date_rows.empty:
        bits = []
        for _, r in date_rows.head(5).iterrows():
            mn = r.get("min")
            mx = r.get("max")
            bits.append(f"`{r['column']}` ({mn} → {mx})")
        lines.append("- **Date/time columns:** " + "; ".join(bits))

    return "\n".join(lines)


def _df_from_table_like(table: List[List[Any]]) -> Optional[pd.DataFrame]:
    if not table:
        return None

    # Convert to DataFrame; drop fully empty rows.
    df = pd.DataFrame(table)
    df = df.dropna(how="all").reset_index(drop=True)
    if df.shape[0] < 2 or df.shape[1] < 2:
        # Need at least header + one row for useful analysis.
        return None

    header = df.iloc[0].astype(str).tolist()
    data = df.iloc[1:].copy()
    data.columns = header

    # Drop columns that are all empty.
    data = data.dropna(axis=1, how="all")
    if data.shape[1] < 2:
        return None

    # Normalize columns and return.
    norm_cols = [_normalize_col_name(c) for c in data.columns]
    norm_cols = _dedupe_columns(norm_cols)
    data.columns = norm_cols

    return data


def _extract_tables_from_pdf(contents: bytes) -> List[pd.DataFrame]:
    import io

    import pdfplumber

    dfs: List[pd.DataFrame] = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            # best-effort: extract_tables returns a list of tables (list-of-rows).
            try:
                tables = page.extract_tables()
            except Exception:
                continue

            for t in tables or []:
                df = _df_from_table_like(t)
                if df is not None:
                    dfs.append(df)
    return dfs


def _extract_tables_from_docx(contents: bytes) -> List[pd.DataFrame]:
    import io

    from docx import Document

    dfs: List[pd.DataFrame] = []
    doc = Document(io.BytesIO(contents))
    for tbl in doc.tables:
        rows: List[List[str]] = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        df = _df_from_table_like(rows)
        if df is not None:
            dfs.append(df)
    return dfs


def _extract_text_from_pdf(contents: bytes) -> str:
    import io

    import pdfplumber

    out: List[str] = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            out.append(txt)
    return "\n".join(out).strip()


def _extract_text_from_docx(contents: bytes) -> str:
    import io

    from docx import Document

    doc = Document(io.BytesIO(contents))
    parts = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
    return "\n".join(parts).strip()


def load_dataset(uploaded_file) -> LoadedDataset:
    """
    Load a dataset from an uploaded file.

    Supported:
    - CSV / XLSX: load into pandas DataFrame
    - PDF / DOCX: best-effort extract tables into DataFrames; otherwise extract text
    """
    filename = (getattr(uploaded_file, "name", "") or "").lower()
    if not filename:
        raise ValueError("Uploaded file is missing a filename.")

    # Streamlit UploadedFile supports read() -> bytes
    contents = uploaded_file.read()
    if not contents:
        raise ValueError("Uploaded file is empty.")

    # Conservative guard to keep app responsive.
    max_bytes = 100 * 1024 * 1024  # 100MB
    if len(contents) > max_bytes:
        raise ValueError("File too large. Please upload a file under 100MB.")

    if filename.endswith(".csv"):
        import io

        df = pd.read_csv(io.BytesIO(contents))
        norm_cols = [_normalize_col_name(c) for c in df.columns]
        df.columns = _dedupe_columns(norm_cols)
        sampled = False
        max_rows = 1_000_000
        if df.shape[0] > max_rows:
            df = df.sample(max_rows, random_state=0)
            sampled = True
        return LoadedDataset(
            mode="dataframe",
            df=df,
            schema=_build_schema(df),
            extracted_text=None,
            summary={**_summarize_dataframe(df), "sampled": sampled},
        )

    if filename.endswith((".xlsx", ".xlsm", ".xls")):
        import io

        # For MVP: accept .xlsx (and tolerate .xlsm/.xlsm by reading with read_excel).
        df = pd.read_excel(io.BytesIO(contents))
        norm_cols = [_normalize_col_name(c) for c in df.columns]
        df.columns = _dedupe_columns(norm_cols)
        sampled = False
        max_rows = 1_000_000
        if df.shape[0] > max_rows:
            df = df.sample(max_rows, random_state=0)
            sampled = True
        return LoadedDataset(
            mode="dataframe",
            df=df,
            schema=_build_schema(df),
            extracted_text=None,
            summary={**_summarize_dataframe(df), "sampled": sampled},
        )

    if filename.endswith(".pdf"):
        tables = _extract_tables_from_pdf(contents)
        if tables:
            # Pick the largest extracted table (by rows * cols).
            tables_sorted = sorted(tables, key=lambda d: d.shape[0] * d.shape[1], reverse=True)
            df = tables_sorted[0]
            return LoadedDataset(
                mode="dataframe",
                df=df,
                schema=_build_schema(df),
                extracted_text=None,
                summary=_summarize_dataframe(df),
                tables=tables_sorted,
                table_index=0,
            )

        extracted_text = _extract_text_from_pdf(contents)
        return LoadedDataset(
            mode="text",
            df=None,
            schema=[],
            extracted_text=extracted_text,
            summary={
                "extraction_mode": "text",
                "chars": len(extracted_text),
            },
        )

    if filename.endswith(".docx"):
        tables = _extract_tables_from_docx(contents)
        if tables:
            tables_sorted = sorted(tables, key=lambda d: d.shape[0] * d.shape[1], reverse=True)
            df = tables_sorted[0]
            return LoadedDataset(
                mode="dataframe",
                df=df,
                schema=_build_schema(df),
                extracted_text=None,
                summary=_summarize_dataframe(df),
                tables=tables_sorted,
                table_index=0,
            )

        extracted_text = _extract_text_from_docx(contents)
        return LoadedDataset(
            mode="text",
            df=None,
            schema=[],
            extracted_text=extracted_text,
            summary={
                "extraction_mode": "text",
                "chars": len(extracted_text),
            },
        )

    raise ValueError("Unsupported file type. Upload CSV, XLSX, PDF, or DOCX.")

